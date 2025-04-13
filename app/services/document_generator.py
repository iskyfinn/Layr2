import os
import json
import re
from typing import Dict, Any, List, Optional, Tuple
import datetime
import uuid
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter, defaultdict
import json

# Try to download nltk resources (might need to handle this differently in production)
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except:
    pass  # Handle silently if we can't download


class EnhancedDocumentGenerator:
    """
    Enhanced service for autonomously generating High Level Design Documents (HLDDs)
    with AI-driven content generation and recommendations.
    """
    
    def __init__(self, output_dir=None):
        """Initialize the document generator"""
        self.output_dir = output_dir
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            
        # Knowledge bases for document generation
        self.patterns_kb = self._load_patterns_knowledge_base()
        self.technologies_kb = self._load_technologies_knowledge_base()
        self.domain_kb = self._load_domain_knowledge_base()
        
        # Track generated documents and their metrics
        self.generated_documents = {}
    
    def generate_hldd(self, 
                     application: Optional[Any] = None, 
                     user_inputs: Optional[Dict] = None, 
                     hldd_data: Optional[Dict] = None,
                     auto_enhance: bool = True) -> Dict[str, Any]:
        """
        Generate a High Level Design Document (HLDD) with intelligent content enhancement
        
        Parameters:
        - application: Application model instance (optional)
        - user_inputs: User-provided inputs (optional)
        - hldd_data: Pre-generated HLDD data structure (optional)
        - auto_enhance: Whether to intelligently enhance content
        
        Returns:
        - Dict with document info including filename
        """
        try:
            # If no hldd_data is provided, generate it using recommendations
            if hldd_data is None:
                if hasattr(self, 'generate_hldd_recommendation'):
                    hldd_data = self.generate_hldd_recommendation(user_inputs or {})
                else:
                    # Try to import from hldd_generator service
                    try:
                        from app.services.hldd_generator import generate_hldd_recommendation
                        
                        # Prepare inputs for recommendation
                        recommendation_inputs = {}
                        if application:
                            recommendation_inputs['name'] = getattr(application, 'name', 'Application')
                            recommendation_inputs['description'] = getattr(application, 'description', '')
                        
                        # Merge with user inputs if provided
                        if user_inputs:
                            recommendation_inputs.update(user_inputs)
                        
                        # Generate recommendation
                        hldd_data = generate_hldd_recommendation(recommendation_inputs)
                    except ImportError:
                        # Fallback to creating basic structure
                        hldd_data = self._create_basic_hldd_structure(application, user_inputs)
            
            # If auto_enhance is enabled, enhance the document content
            if auto_enhance:
                hldd_data = self._enhance_document_content(hldd_data)
            
            # Extract basic application info
            basic_info = hldd_data.get('basicInfo', {})
            app_name = basic_info.get('name', 'Application')
            
            # Generate unique ID for tracking
            document_id = str(uuid.uuid4())
            
            # Generate filename
            date_str = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            file_name = f"{app_name.replace(' ', '_')}_HLDD_{date_str}.docx"
            
            # Determine file path
            if self.output_dir:
                file_path = os.path.join(self.output_dir, file_name)
            else:
                file_path = file_name
            
            # Generate document
            doc = self._create_document(
                basic_info, 
                hldd_data.get('technologyStack', {}), 
                hldd_data.get('architectureInfo', {})
            )
            
            # Save document
            doc.save(file_path)
            
            # Track the document
            self.generated_documents[document_id] = {
                'document_id': document_id,
                'file_name': file_name,
                'file_path': file_path,
                'creation_time': datetime.datetime.now().isoformat(),
                'app_name': app_name,
                'content_metrics': self._calculate_document_metrics(hldd_data)
            }
            
            return {
                'document_id': document_id,
                'file_name': file_name,
                'file_path': file_path,
                'success': True,
                'hldd_data': hldd_data,  # Return the generated HLDD data
                'content_metrics': self.generated_documents[document_id]['content_metrics']
            }
        
        except Exception as e:
            import traceback
            return {
                'error': str(e),
                'traceback': traceback.format_exc(),
                'success': False
            }

    def get_document_quality_feedback(self, document_id: str) -> Dict[str, Any]:
        """
        Get quality feedback for a generated document
        
        Parameters:
        - document_id: ID of the previously generated document
        
        Returns:
        - Dict with quality metrics and improvement suggestions
        """
        if document_id not in self.generated_documents:
            return {
                'error': 'Document not found',
                'success': False
            }
        
        document_info = self.generated_documents[document_id]
        
        # Extract quality metrics
        metrics = document_info['content_metrics']
        
        # Generate improvement suggestions
        suggestions = []
        
        if metrics['completeness'] < 0.8:
            suggestions.append({
                'area': 'Content Completeness',
                'suggestion': 'Add more details to sections that are sparse',
                'sections': metrics.get('incomplete_sections', [])
            })
        
        if metrics['technical_depth'] < 0.7:
            suggestions.append({
                'area': 'Technical Depth',
                'suggestion': 'Enhance technical details, especially regarding architecture patterns and technology choices',
                'focus_areas': ['Architecture Overview', 'Technology Stack']
            })
        
        if metrics['clarity'] < 0.7:
            suggestions.append({
                'area': 'Clarity and Readability',
                'suggestion': 'Improve clarity by breaking down complex sentences and adding examples',
                'sections': metrics.get('complex_sections', [])
            })
        
        # Overall rating
        overall_quality = (metrics['completeness'] + metrics['technical_depth'] + metrics['clarity']) / 3
        quality_rating = 'Excellent' if overall_quality > 0.9 else 'Good' if overall_quality > 0.7 else 'Needs Improvement'
        
        return {
            'document_id': document_id,
            'quality_metrics': metrics,
            'quality_rating': quality_rating,
            'improvement_suggestions': suggestions,
            'success': True
        }
    
    def enhance_existing_document(self, file_path: str, enhancement_options: Dict[str, bool] = None) -> Dict[str, Any]:
        """
        Enhance an existing HLDD document
        
        Parameters:
        - file_path: Path to existing document
        - enhancement_options: Options for what to enhance
        
        Returns:
        - Dict with enhanced document info
        """
        try:
            # Set default enhancement options if not provided
            if enhancement_options is None:
                enhancement_options = {
                    'expand_architecture': True,
                    'add_security_details': True,
                    'improve_diagrams': True,
                    'update_technology': True,
                    'add_best_practices': True
                }
            
            # Load existing document
            doc = Document(file_path)
            
            # Extract document content
            extracted_content = self._extract_document_content(doc)
            
            # Enhance the content
            enhanced_content = self._enhance_extracted_content(
                extracted_content,
                enhancement_options
            )
            
            # Generate new document with enhanced content
            new_doc = self._create_document_from_content(enhanced_content)
            
            # Save enhanced document
            base_name, ext = os.path.splitext(file_path)
            enhanced_file_path = f"{base_name}_enhanced{ext}"
            
            if self.output_dir:
                enhanced_file_path = os.path.join(self.output_dir, os.path.basename(enhanced_file_path))
            
            new_doc.save(enhanced_file_path)
            
            # Generate document ID for tracking
            document_id = str(uuid.uuid4())
            
            # Track the enhanced document
            self.generated_documents[document_id] = {
                'document_id': document_id,
                'file_name': os.path.basename(enhanced_file_path),
                'file_path': enhanced_file_path,
                'creation_time': datetime.datetime.now().isoformat(),
                'app_name': extracted_content.get('app_name', 'Enhanced Application'),
                'is_enhanced': True,
                'original_document': file_path,
                'content_metrics': self._calculate_document_metrics(enhanced_content)
            }
            
            return {
                'document_id': document_id,
                'file_name': os.path.basename(enhanced_file_path),
                'file_path': enhanced_file_path,
                'success': True,
                'enhancement_summary': self._generate_enhancement_summary(
                    extracted_content,
                    enhanced_content
                )
            }
        
        except Exception as e:
            import traceback
            return {
                'error': str(e),
                'traceback': traceback.format_exc(),
                'success': False
            }
    
    def generate_document_sections(self, 
                                  section_type: str, 
                                  application_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate specific document sections autonomously
        
        Parameters:
        - section_type: Type of section to generate ('architecture', 'security', 'technology', etc.)
        - application_info: Application information to use for generation
        
        Returns:
        - Dict with generated section content
        """
        try:
            # Map section types to generation functions
            section_generators = {
                'architecture': self._generate_architecture_section,
                'security': self._generate_security_section,
                'technology': self._generate_technology_section,
                'data': self._generate_data_section,
                'deployment': self._generate_deployment_section,
                'operations': self._generate_operations_section,
                'introduction': self._generate_introduction_section,
                'appendix': self._generate_appendix_section
            }
            
            # Check if requested section type is supported
            if section_type not in section_generators:
                return {
                    'error': f"Unsupported section type: {section_type}",
                    'supported_types': list(section_generators.keys()),
                    'success': False
                }
            
            # Generate the section
            section_content = section_generators[section_type](application_info)
            
            return {
                'section_type': section_type,
                'content': section_content,
                'success': True
            }
        
        except Exception as e:
            return {
                'error': str(e),
                'section_type': section_type,
                'success': False
            }
    
    def suggest_content_improvements(self, hldd_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Suggest improvements to HLDD content
        
        Parameters:
        - hldd_data: Current HLDD data structure
        
        Returns:
        - Dict with improvement suggestions
        """
        try:
            # Calculate current metrics
            metrics = self._calculate_document_metrics(hldd_data)
            
            # Generate improvement suggestions
            suggestions = {
                'general': [],
                'sections': {}
            }
            
            # General suggestions based on metrics
            if metrics['completeness'] < 0.8:
                suggestions['general'].append(
                    "Increase overall content completeness by expanding sections with minimal content"
                )
            
            if metrics['technical_depth'] < 0.7:
                suggestions['general'].append(
                    "Add more technical details about architecture patterns and implementation approaches"
                )
            
            if metrics['clarity'] < 0.7:
                suggestions['general'].append(
                    "Improve clarity by simplifying complex sections and adding explanatory content"
                )
            
            # Analyze individual sections
            architecture_info = hldd_data.get('architectureInfo', {})
            if 'components' not in architecture_info or len(architecture_info.get('components', [])) < 3:
                suggestions['sections']['architecture'] = [
                    "Add more details about system components and their interactions",
                    "Include architectural patterns and their implementation details",
                    "Describe component responsibilities and interfaces"
                ]
            
            tech_stack = hldd_data.get('technologyStack', {})
            if 'categories' not in tech_stack or len(tech_stack.get('categories', [])) < 2:
                suggestions['sections']['technology'] = [
                    "Expand technology stack with more specific categories",
                    "Include version information and rationale for technology choices",
                    "Add details about how technologies integrate with each other"
                ]
            
            # Check for security information
            if 'security' not in architecture_info or not architecture_info.get('security'):
                suggestions['sections']['security'] = [
                    "Add security architecture details",
                    "Include authentication and authorization approaches",
                    "Describe data protection mechanisms"
                ]
            
            return {
                'current_metrics': metrics,
                'improvement_suggestions': suggestions,
                'success': True
            }
        
        except Exception as e:
            return {
                'error': str(e),
                'success': False
            }
    
    # Internal methods
    
    def _create_basic_hldd_structure(self, application, user_inputs):
        """Create a basic HLDD structure when no generator service is available"""
        # Extract basic application info
        app_name = getattr(application, 'name', 'Application') if application else user_inputs.get('name', 'Application')
        description = getattr(application, 'description', '') if application else user_inputs.get('description', '')
        
        # Build basic structure
        return {
            'basicInfo': {
                'name': app_name,
                'description': description,
                'version': '1.0',
                'author': 'Architecture Team'
            },
            'technologyStack': {
                'overview': 'Technology stack for the application',
                'categories': [
                    {
                        'name': 'Frontend',
                        'technologies': [
                            {
                                'name': 'React',
                                'version': 'Latest',
                                'purpose': 'User interface framework'
                            }
                        ]
                    },
                    {
                        'name': 'Backend',
                        'technologies': [
                            {
                                'name': 'Python (Flask)',
                                'version': 'Latest',
                                'purpose': 'Server-side logic and API'
                            }
                        ]
                    }
                ]
            },
            'architectureInfo': {
                'overview': 'System architecture overview',
                'architectural_style': 'Layered',
                'components': [
                    {
                        'name': 'Web Layer',
                        'description': 'User interface and presentation layer'
                    },
                    {
                        'name': 'Application Layer',
                        'description': 'Business logic and application services'
                    },
                    {
                        'name': 'Data Layer',
                        'description': 'Data access and storage'
                    }
                ]
            }
        }
    
    def _enhance_document_content(self, hldd_data):
        """Enhance document content with AI-generated suggestions"""
        enhanced_data = hldd_data.copy()
        
        # Enhance architecture information
        if 'architectureInfo' in enhanced_data:
            arch_info = enhanced_data['architectureInfo']
            
            # Add principles if missing
            if 'principles' not in arch_info or not arch_info['principles']:
                arch_info['principles'] = self._generate_architecture_principles(
                    enhanced_data.get('basicInfo', {}).get('description', ''),
                    arch_info.get('architectural_style', '')
                )
            
            # Enhance component descriptions
            if 'components' in arch_info:
                arch_info['components'] = self._enhance_component_descriptions(
                    arch_info['components'],
                    arch_info.get('architectural_style', '')
                )
            
            # Add security section if missing
            if 'security' not in arch_info or not arch_info['security']:
                arch_info['security'] = self._generate_security_recommendations(
                    enhanced_data.get('basicInfo', {}).get('description', '')
                )
            
            # Add deployment section if missing
            if 'deployment' not in arch_info or not arch_info['deployment']:
                arch_info['deployment'] = self._generate_deployment_recommendations(
                    arch_info.get('architectural_style', ''),
                    enhanced_data.get('technologyStack', {})
                )
            
            # Add operations section if missing
            if 'operations' not in arch_info or not arch_info['operations']:
                arch_info['operations'] = self._generate_operations_recommendations(
                    arch_info.get('architectural_style', ''),
                    enhanced_data.get('technologyStack', {})
                )
        
        # Enhance technology stack
        if 'technologyStack' in enhanced_data:
            tech_stack = enhanced_data['technologyStack']
            
            # Add technology rationale if missing
            for category in tech_stack.get('categories', []):
                for tech in category.get('technologies', []):
                    if 'rationale' not in tech:
                        tech['rationale'] = self._generate_technology_rationale(
                            tech['name'],
                            tech['purpose'],
                            category['name']
                        )
        
        # Add data architecture if missing
        if 'dataArchitecture' not in enhanced_data:
            enhanced_data['dataArchitecture'] = self._generate_data_architecture(
                enhanced_data.get('basicInfo', {}).get('description', ''),
                enhanced_data.get('architectureInfo', {}).get('architectural_style', ''),
                enhanced_data.get('technologyStack', {})
            )
        
        return enhanced_data
    
    def _calculate_document_metrics(self, hldd_data):
        """Calculate quality metrics for the document content"""
        metrics = {
            'completeness': 0.0,
            'technical_depth': 0.0,
            'clarity': 0.0,
            'incomplete_sections': [],
            'complex_sections': []
        }
        
        # Calculate completeness
        required_sections = [
            'basicInfo', 'architectureInfo', 'technologyStack'
        ]
        
        required_arch_sections = [
            'overview', 'architectural_style', 'components', 'principles'
        ]
        
        required_tech_sections = [
            'overview', 'categories'
        ]
        
        # Check main sections
        main_sections_present = sum(1 for section in required_sections if section in hldd_data)
        main_completeness = main_sections_present / len(required_sections)
        
        # Check architecture sections
        arch_info = hldd_data.get('architectureInfo', {})
        arch_sections_present = sum(1 for section in required_arch_sections if section in arch_info)
        arch_completeness = arch_sections_present / len(required_arch_sections) if arch_sections_present > 0 else 0
        
        # Check technology sections
        tech_stack = hldd_data.get('technologyStack', {})
        tech_sections_present = sum(1 for section in required_tech_sections if section in tech_stack)
        tech_completeness = tech_sections_present / len(required_tech_sections) if tech_sections_present > 0 else 0
        
        # Identify incomplete sections
        if 'architectureInfo' not in hldd_data or not hldd_data.get('architectureInfo'):
            metrics['incomplete_sections'].append('Architecture Information')
        elif arch_completeness < 0.7:
            metrics['incomplete_sections'].append('Architecture Information')
        
        if 'technologyStack' not in hldd_data or not hldd_data.get('technologyStack'):
            metrics['incomplete_sections'].append('Technology Stack')
        elif tech_completeness < 0.7:
            metrics['incomplete_sections'].append('Technology Stack')
        
        # Calculate overall completeness
        metrics['completeness'] = (main_completeness + arch_completeness + tech_completeness) / 3
        
        # Calculate technical depth
        tech_terms_count = 0
        total_text_length = 0
        
        # Count technical terms in architecture
        if 'architectureInfo' in hldd_data:
            arch_info = hldd_data['architectureInfo']
            
            # Count in overview
            overview_text = arch_info.get('overview', '')
            tech_terms_count += self._count_technical_terms(overview_text)
            total_text_length += len(overview_text.split())
            
            # Count in components
            for component in arch_info.get('components', []):
                component_desc = component.get('description', '')
                tech_terms_count += self._count_technical_terms(component_desc)
                total_text_length += len(component_desc.split())
        
        # Count technical terms in technology stack
        if 'technologyStack' in hldd_data:
            tech_stack = hldd_data['technologyStack']
            
            # Count in overview
            overview_text = tech_stack.get('overview', '')
            tech_terms_count += self._count_technical_terms(overview_text)
            total_text_length += len(overview_text.split())
            
            # Count in categories
            for category in tech_stack.get('categories', []):
                for tech in category.get('technologies', []):
                    purpose_text = tech.get('purpose', '')
                    tech_terms_count += self._count_technical_terms(purpose_text)
                    total_text_length += len(purpose_text.split())
                    
                    if 'rationale' in tech:
                        rationale_text = tech['rationale']
                        tech_terms_count += self._count_technical_terms(rationale_text)
                        total_text_length += len(rationale_text.split())
        
        # Calculate technical depth score
        if total_text_length > 0:
            technical_term_ratio = tech_terms_count / total_text_length
            metrics['technical_depth'] = min(1.0, technical_term_ratio * 5)  # Scale ratio appropriately
        else:
            metrics['technical_depth'] = 0.0
        
        # Calculate clarity
        complex_sentences = 0
        total_sentences = 0
        
        # Check clarity in architecture
        if 'architectureInfo' in hldd_data:
            arch_info = hldd_data['architectureInfo']
            
            # Check overview
            overview_text = arch_info.get('overview', '')
            sentences = sent_tokenize(overview_text) if overview_text else []
            total_sentences += len(sentences)
            complex_sentences += sum(1 for s in sentences if self._is_complex_sentence(s))
            
            # Check if architecture overview is complex
            if overview_text and self._is_section_complex(overview_text):
                metrics['complex_sections'].append('Architecture Overview')
            
            # Check components
            for component in arch_info.get('components', []):
                component_desc = component.get('description', '')
                sentences = sent_tokenize(component_desc) if component_desc else []
                total_sentences += len(sentences)
                complex_sentences += sum(1 for s in sentences if self._is_complex_sentence(s))
        
        # Check clarity in technology stack
        if 'technologyStack' in hldd_data:
            tech_stack = hldd_data['technologyStack']
            
            # Check overview
            overview_text = tech_stack.get('overview', '')
            sentences = sent_tokenize(overview_text) if overview_text else []
            total_sentences += len(sentences)
            complex_sentences += sum(1 for s in sentences if self._is_complex_sentence(s))
            
            # Check if technology overview is complex
            if overview_text and self._is_section_complex(overview_text):
                metrics['complex_sections'].append('Technology Stack Overview')
        
        # Calculate clarity score
        if total_sentences > 0:
            clear_sentence_ratio = 1 - (complex_sentences / total_sentences)
            metrics['clarity'] = clear_sentence_ratio
        else:
            metrics['clarity'] = 0.5  # Default mid-range score if no sentences to analyze
        
        return metrics
    
    def _count_technical_terms(self, text):
        """Count technical terms in text"""
        if not text:
            return 0
        
        # This would be a set of technical terms in a real implementation
        technical_terms = set([
            'api', 'rest', 'http', 'json', 'xml', 'database', 'sql', 'nosql', 
            'architecture', 'microservice', 'serverless', 'container', 'docker',
            'kubernetes', 'cloud', 'aws', 'azure', 'authentication', 'authorization',
            'scaling', 'cache', 'load balancer', 'message queue', 'security',
            'encryption', 'framework', 'react', 'angular', 'vue', 'flask', 'django',
            'express', 'spring', 'postgresql', 'mysql', 'mongodb', 'algorithm',
            'protocol', 'ssl', 'tls', 'vpc', 'subnet', 'firewall', 'cicd', 'git',
            'deployment', 'monitoring', 'logging', 'lambda', 'function', 'service'
        ])
        
        # Tokenize and count terms
        words = word_tokenize(text.lower())
        return sum(1 for word in words if word in technical_terms)
    
    def _is_complex_sentence(self, sentence):
        """Check if a sentence is complex"""
        # Simple heuristics for complexity
        word_count = len(word_tokenize(sentence))
        return word_count > 25  # Long sentences tend to be complex
    
    def _is_section_complex(self, text):
        """Check if a section is complex based on its text"""
        # Simple heuristics for section complexity
        sentences = sent_tokenize(text)
        complex_sentence_count = sum(1 for s in sentences if self._is_complex_sentence(s))
        return complex_sentence_count / len(sentences) > 0.3 if sentences else False
    
    def _generate_architecture_principles(self, description, arch_style):
        """Generate architecture principles based on description and style"""
        # This would use more advanced NLP in a real implementation
        principles = [
            f"Scalability: The system should scale horizontally following {arch_style} patterns",
            f"Reliability: The system should be resilient to failures and maintain high availability",
            f"Security: The system should implement defense in depth and follow security best practices",
            f"Maintainability: The system should be easy to maintain and update"
        ]
        
        # Add principles based on keywords in description
        description_lower = description.lower()
        
        if any(term in description_lower for term in ['performance', 'fast', 'speed', 'quick']):
            principles.append("Performance: The system should optimize for high performance and low latency")
        
        if any(term in description_lower for term in ['cloud', 'aws', 'azure', 'google']):
            principles.append("Cloud-Native: The system should leverage cloud services and follow cloud design principles")
        
        if any(term in description_lower for term in ['compliance', 'regulation', 'gdpr', 'hipaa', 'pci']):
            principles.append("Compliance: The system should adhere to relevant regulatory requirements")
        
        return principles
    
    def _enhance_component_descriptions(self, components, arch_style):
        """Enhance component descriptions with more details"""
        enhanced_components = []
        
        for component in components:
            name = component.get('name', '')
            description = component.get('description', '')
            
            # Enhance description if it's minimal
            if len(description.split()) < 10:
                enhanced_description = self._generate_enhanced_component_description(
                    name, description, arch_style
                )
                component['description'] = enhanced_description
            
            # Add responsibilities if missing
            if 'responsibilities' not in component:
                component['responsibilities'] = self._generate_component_responsibilities(
                    name, description, arch_style
                )
            
            enhanced_components.append(component)
        
        return enhanced_components
    
    def _generate_enhanced_component_description(self, name, description, arch_style):
        """Generate an enhanced description for a component"""
        # This would use more advanced NLP in a real implementation
        
        # Check if this is a known component type
        lower_name = name.lower()
        
        if 'api' in lower_name or 'gateway' in lower_name:
            return (
                f"The {name} serves as the entry point for client requests, "
                f"providing routing, transformation, and security controls for API calls. "
                f"It implements {arch_style} principles for API management and integrates "
                f"with backend services."
            )
        
        elif 'frontend' in lower_name or 'ui' in lower_name:
            return (
                f"The {name} provides the user interface layer of the application, "
                f"handling presentation logic and user interactions. It follows "
                f"responsive design principles and communicates with backend services "
                f"through API calls."
            )
        
        elif 'service' in lower_name or 'backend' in lower_name:
            return (
                f"The {name} implements core business logic and processing capabilities, "
                f"following {arch_style} architecture principles. It handles data processing, "
                f"business rules, and integration with other system components."
            )
        
        elif 'data' in lower_name or 'database' in lower_name:
            return (
                f"The {name} manages data storage, retrieval, and persistence, "
                f"providing data access services to other components. It ensures "
                f"data integrity, consistency, and security while optimizing "
                f"for performance and scalability."
            )
        
        elif 'auth' in lower_name or 'security' in lower_name:
                    return (
                        f"The {name} handles authentication, authorization, and security controls "
                        f"for the application. It integrates with identity providers, manages access "
                        f"controls, and implements security policies to protect system resources."
                    )
        
        elif 'monitor' in lower_name or 'logging' in lower_name:
            return (
                f"The {name} provides monitoring, logging, and observability capabilities "
                f"for the system. It collects metrics, logs, and traces to enable operational "
                f"visibility, troubleshooting, and performance analysis."
            )
        
        # Default enhancement if no specific type identified
        return (
            f"The {name} is a key component in the {arch_style} architecture, "
            f"providing essential functionality for the system. {description} "
            f"It interacts with other components through well-defined interfaces "
            f"and follows system-wide design principles."
        )
    
    def _generate_component_responsibilities(self, name, description, arch_style):
        """Generate responsibilities for a component based on its name and description"""
        # This would use more advanced NLP in a real implementation
        lower_name = name.lower()
        
        # Define responsibilities based on component type
        if 'api' in lower_name or 'gateway' in lower_name:
            return [
                "Route and direct client requests to appropriate services",
                "Implement authentication and authorization checks",
                "Transform requests and responses as needed",
                "Provide rate limiting and throttling",
                "Handle API versioning and backward compatibility"
            ]
        
        elif 'frontend' in lower_name or 'ui' in lower_name:
            return [
                "Present information to users through intuitive interfaces",
                "Handle user interactions and input validation",
                "Communicate with backend services via API calls",
                "Manage client-side state and navigation",
                "Implement responsive design for multiple device types"
            ]
        
        elif 'service' in lower_name or 'backend' in lower_name:
            return [
                "Implement business logic and domain rules",
                "Process data and perform calculations",
                "Orchestrate operations across multiple components",
                "Manage transactional integrity and consistency",
                "Provide service interfaces to other components"
            ]
        
        elif 'data' in lower_name or 'database' in lower_name:
            return [
                "Store and retrieve data efficiently",
                "Ensure data consistency and integrity",
                "Implement data access patterns and optimization",
                "Manage database schema and migrations",
                "Support backup and recovery operations"
            ]
        
        elif 'auth' in lower_name or 'security' in lower_name:
            return [
                "Authenticate users and verify identities",
                "Authorize access to resources based on permissions",
                "Manage user sessions and tokens",
                "Implement security controls and policies",
                "Audit security events and access attempts"
            ]
        
        elif 'monitor' in lower_name or 'logging' in lower_name:
            return [
                "Collect system metrics and performance data",
                "Aggregate and store logs from various components",
                "Provide alerting for abnormal conditions",
                "Support troubleshooting and root cause analysis",
                "Generate operational reports and dashboards"
            ]
        
        # Generic responsibilities for unknown component types
        return [
            f"Provide core functionality within the {arch_style} architecture",
            "Interact with other system components through defined interfaces",
            "Maintain performance and reliability for assigned functions",
            "Implement error handling and resilience mechanisms",
            "Support system-wide design principles and standards"
        ]
    
    def _generate_security_recommendations(self, description):
        """Generate security recommendations based on description"""
        # This would use more advanced NLP in a real implementation
        description_lower = description.lower()
        
        # Determine security sensitivity
        security_level = 'medium'
        
        high_sensitivity_terms = [
            'financial', 'banking', 'payment', 'healthcare', 'medical', 'personal',
            'pii', 'phi', 'hipaa', 'gdpr', 'pci', 'classified', 'confidential',
            'sensitive', 'security', 'compliance', 'regulation'
        ]
        
        low_sensitivity_terms = [
            'public', 'open', 'informational', 'educational', 'content',
            'blog', 'marketing', 'static'
        ]
        
        # Check for high sensitivity indicators
        if any(term in description_lower for term in high_sensitivity_terms):
            security_level = 'high'
        # Check for low sensitivity indicators
        elif any(term in description_lower for term in low_sensitivity_terms):
            security_level = 'low'
        
        # Generate security recommendations based on sensitivity
        if security_level == 'high':
            return {
                'overview': 'Comprehensive security architecture for highly sensitive data',
                'authentication': 'Multi-factor authentication with biometric options',
                'authorization': 'Fine-grained role-based access control (RBAC) with principle of least privilege',
                'data_protection': 'End-to-end encryption for all data in transit and at rest, with key rotation',
                'network_security': 'Zero-trust architecture with micro-segmentation and continuous monitoring',
                'compliance': 'Regular compliance audits for relevant regulations (HIPAA, GDPR, PCI DSS, etc.)',
                'security_testing': 'Comprehensive security testing including penetration testing and code security analysis'
            }
        elif security_level == 'medium':
            return {
                'overview': 'Enhanced security measures for business applications',
                'authentication': 'Strong password policies with multi-factor authentication',
                'authorization': 'Role-based access control with appropriate separation of duties',
                'data_protection': 'Encryption for sensitive data in transit and at rest',
                'network_security': 'Network segmentation, firewalls, and intrusion detection',
                'compliance': 'Security controls aligned with industry best practices',
                'security_testing': 'Regular security testing and vulnerability scanning'
            }
        else:  # low
            return {
                'overview': 'Standard security measures for low-sensitivity applications',
                'authentication': 'Username/password authentication with secure password policies',
                'authorization': 'Basic access control for resource protection',
                'data_protection': 'TLS for data in transit and standard security for data at rest',
                'network_security': 'Firewall protection and standard security monitoring',
                'compliance': 'Security controls based on organizational standards',
                'security_testing': 'Regular vulnerability scanning and updates'
            }
    
    def _generate_deployment_recommendations(self, arch_style, tech_stack):
        """Generate deployment recommendations based on architecture style and technology stack"""
        # This would use more advanced NLP in a real implementation
        
        # Determine deployment approach based on architecture style
        deployment_approach = 'containers'  # default
        
        if arch_style.lower() == 'serverless':
            deployment_approach = 'serverless'
        elif arch_style.lower() in ['monolithic', 'layered']:
            deployment_approach = 'virtual_machines'
        elif arch_style.lower() in ['microservices', 'service-oriented']:
            deployment_approach = 'containers'
        
        # Generate environments
        environments = [
            {'name': 'Development', 'purpose': 'Used for development and testing by the development team'},
            {'name': 'QA/Test', 'purpose': 'Used for quality assurance and integration testing'},
            {'name': 'Staging', 'purpose': 'Production-like environment for final validation before deployment'},
            {'name': 'Production', 'purpose': 'Live environment serving end users'}
        ]
        
        # Generate deployment recommendations based on approach
        if deployment_approach == 'serverless':
            return {
                'overview': 'Serverless deployment model leveraging cloud provider managed services',
                'environments': environments,
                'infrastructure': 'Cloud-native serverless architecture with Function-as-a-Service (FaaS)',
                'deployment_process': 'Automated CI/CD pipeline with infrastructure as code',
                'scaling': 'Automatic scaling based on demand, managed by cloud provider',
                'considerations': [
                    'Cold start latency for infrequently used functions',
                    'Monitoring and observability across distributed services',
                    'Cost management for high-volume function invocations',
                    'Deployment packaging and dependency management'
                ]
            }
        elif deployment_approach == 'containers':
            return {
                'overview': 'Container-based deployment using orchestration for scalability and management',
                'environments': environments,
                'infrastructure': 'Container orchestration platform (Kubernetes or equivalent)',
                'deployment_process': 'CI/CD pipeline with container registry and automated deployments',
                'scaling': 'Horizontal scaling of containers based on resource utilization',
                'considerations': [
                    'Container resource allocation and limits',
                    'Service discovery and network configuration',
                    'Persistent storage management',
                    'Health checks and self-healing capabilities',
                    'Image security and vulnerability scanning'
                ]
            }
        else:  # virtual_machines
            return {
                'overview': 'Virtual machine deployment with automation for consistency and reliability',
                'environments': environments,
                'infrastructure': 'Cloud or on-premises virtual machines with load balancing',
                'deployment_process': 'CI/CD pipeline with infrastructure automation',
                'scaling': 'Horizontal and vertical scaling based on demand',
                'considerations': [
                    'VM image management and consistency',
                    'Resource provisioning and capacity planning',
                    'Backup and disaster recovery',
                    'Operating system maintenance and updates',
                    'Cost optimization for compute resources'
                ]
            }
    
    def _generate_operations_recommendations(self, arch_style, tech_stack):
        """Generate operations recommendations based on architecture style and technology stack"""
        # This would use more advanced NLP in a real implementation
        
        # Basic operations recommendations for all systems
        operations = {
            'overview': 'Operational considerations for system reliability and maintainability',
            'monitoring': {
                'approach': 'Comprehensive monitoring of system health, performance, and user experience',
                'tools': ['Prometheus', 'Grafana', 'ELK Stack (Elasticsearch, Logstash, Kibana)'],
                'metrics': ['System uptime', 'Response time', 'Error rates', 'Resource utilization', 'Business KPIs']
            },
            'logging': {
                'approach': 'Centralized logging with structured log format',
                'retention': '30 days for standard logs, 1 year for security and audit logs',
                'levels': ['DEBUG', 'INFO', 'WARN', 'ERROR', 'FATAL']
            },
            'backup': {
                'approach': 'Regular automated backups with verification',
                'frequency': 'Daily full backups, hourly incremental backups',
                'retention': '30 days of daily backups, 1 year of monthly backups'
            },
            'disaster_recovery': {
                'approach': 'Documented disaster recovery plan with regular testing',
                'rpo': '1 hour (Recovery Point Objective)',
                'rto': '4 hours (Recovery Time Objective)'
            }
        }
        
        # Enhance with architecture-specific recommendations
        arch_style_lower = arch_style.lower()
        
        if 'microservices' in arch_style_lower or 'service-oriented' in arch_style_lower:
            operations['service_health'] = {
                'approach': 'Service-level health monitoring and circuit breaking',
                'tools': ['Health check APIs', 'Service meshes', 'Distributed tracing'],
                'considerations': [
                    'Service discovery and registration',
                    'Dependency tracking between services',
                    'Distributed transaction monitoring'
                ]
            }
        
        if 'serverless' in arch_style_lower:
            operations['monitoring']['considerations'] = [
                'Function invocation metrics and cold start monitoring',
                'Cost tracking per function',
                'Timeout and error handling',
                'Asynchronous processing monitoring'
            ]
        
        if 'event-driven' in arch_style_lower:
            operations['event_management'] = {
                'approach': 'Monitoring of event flows and queues',
                'tools': ['Queue depth monitoring', 'Event replay capabilities', 'Dead letter queues'],
                'considerations': [
                    'Event ordering and idempotency',
                    'Queue backlogs and processing delays',
                    'Event schema evolution'
                ]
            }
        
        return operations
    
    def _generate_data_architecture(self, description, arch_style, tech_stack):
        """Generate data architecture section based on description and architecture style"""
        # This would use more advanced NLP in a real implementation
        
        # Determine data persistence approach
        data_approach = 'relational'  # default
        description_lower = description.lower()
        arch_style_lower = arch_style.lower()
        
        # Check for indicators of different data approaches
        if any(term in description_lower for term in ['nosql', 'document', 'schema-less', 'flexible schema']):
            data_approach = 'document'
        elif any(term in description_lower for term in ['big data', 'analytics', 'data warehouse', 'reporting']):
            data_approach = 'analytical'
        elif any(term in arch_style_lower for term in ['event', 'streaming', 'message']):
            data_approach = 'event-sourcing'
        
        # Generate data stores based on approach
        if data_approach == 'relational':
            data_stores = [
                {'name': 'Primary Database', 'type': 'Relational (PostgreSQL)', 'purpose': 'Main application data storage'},
                {'name': 'Cache', 'type': 'In-memory (Redis)', 'purpose': 'Performance optimization and session storage'}
            ]
        elif data_approach == 'document':
            data_stores = [
                {'name': 'Document Store', 'type': 'NoSQL (MongoDB)', 'purpose': 'Flexible document storage for application data'},
                {'name': 'Cache', 'type': 'In-memory (Redis)', 'purpose': 'Performance optimization and session storage'}
            ]
        elif data_approach == 'analytical':
            data_stores = [
                {'name': 'Operational Database', 'type': 'Relational (PostgreSQL)', 'purpose': 'Transactional data storage'},
                {'name': 'Data Warehouse', 'type': 'Analytical (Snowflake/BigQuery)', 'purpose': 'Reporting and analytics'},
                {'name': 'Data Lake', 'type': 'Object Storage', 'purpose': 'Raw data storage for processing'}
            ]
        else:  # event-sourcing
            data_stores = [
                {'name': 'Event Store', 'type': 'Append-only log (Kafka)', 'purpose': 'Immutable log of all events'},
                {'name': 'Read Model', 'type': 'Relational/NoSQL', 'purpose': 'Optimized view of data for querying'},
                {'name': 'Cache', 'type': 'In-memory (Redis)', 'purpose': 'Performance optimization'}
            ]
        
        # Generate data models based on description
        # Simple placeholder - in a real implementation this would use NLP to identify entities
        data_models = [
            {'name': 'User', 'attributes': ['id', 'username', 'email', 'profile']},
            {'name': 'Content', 'attributes': ['id', 'title', 'body', 'metadata']}
        ]
        
        # Generate data flows
        data_flows = [
            {'source': 'Client', 'destination': 'API Layer', 'description': 'User requests and responses'},
            {'source': 'API Layer', 'destination': 'Application Services', 'description': 'Processed requests'},
            {'source': 'Application Services', 'destination': 'Data Access Layer', 'description': 'Data operations'},
            {'source': 'Data Access Layer', 'destination': 'Database', 'description': 'Persistence operations'}
        ]
        
        # Add approach-specific flows
        if data_approach == 'analytical':
            data_flows.extend([
                {'source': 'Operational Database', 'destination': 'ETL Process', 'description': 'Data extraction'},
                {'source': 'ETL Process', 'destination': 'Data Warehouse', 'description': 'Transformed data loading'}
            ])
        elif data_approach == 'event-sourcing':
            data_flows.extend([
                {'source': 'Application Services', 'destination': 'Event Store', 'description': 'Event publishing'},
                {'source': 'Event Store', 'destination': 'Event Handlers', 'description': 'Event processing'},
                {'source': 'Event Handlers', 'destination': 'Read Model', 'description': 'View model updates'}
            ])
        
        return {
            'overview': f"Data architecture based on {data_approach} approach, supporting the {arch_style} architecture",
            'data_stores': data_stores,
            'data_models': data_models,
            'data_flows': data_flows,
            'data_management': {
                'governance': 'Defined data ownership and quality standards',
                'security': 'Data encryption and access controls',
                'lifecycle': 'Data retention and archiving policies'
            }
        }
    
    def _generate_technology_rationale(self, tech_name, purpose, category):
        """Generate rationale for a technology choice"""
        # This would use more advanced NLP in a real implementation
        
        # General benefits for common technologies
        tech_name_lower = tech_name.lower()
        
        if 'react' in tech_name_lower:
            return "React provides a component-based architecture that improves code reusability and maintenance. Its virtual DOM implementation ensures efficient UI updates and optimal performance."
        
        elif 'angular' in tech_name_lower:
            return "Angular provides a comprehensive framework with built-in features for forms, routing, and state management, enabling consistent development patterns and reducing the need for third-party libraries."
        
        elif 'vue' in tech_name_lower:
            return "Vue.js offers a progressive framework that can be adopted incrementally, with excellent performance and a gentle learning curve, making it ideal for both simple and complex applications."
        
        elif 'node' in tech_name_lower or 'express' in tech_name_lower:
            return "Node.js with Express provides a lightweight, high-performance server environment with JavaScript across the stack, enabling code sharing and efficient handling of concurrent connections."
        
        elif 'python' in tech_name_lower:
            if 'django' in tech_name_lower:
                return "Django follows the 'batteries-included' philosophy with built-in admin, ORM, and security features, accelerating development while maintaining robustness and security."
            elif 'flask' in tech_name_lower:
                return "Flask provides a lightweight and flexible framework that can be tailored precisely to the application's needs, with excellent extensibility and minimal overhead."
            else:
                return "Python offers excellent readability, extensive libraries, and strong integration capabilities, making it ideal for rapid development and maintenance."
        
        elif 'java' in tech_name_lower or 'spring' in tech_name_lower:
            return "Java with Spring provides enterprise-grade reliability, performance, and security, with a mature ecosystem of libraries and tools supporting complex business requirements."
        
        elif 'postgresql' in tech_name_lower or 'mysql' in tech_name_lower:
            return "This relational database provides ACID compliance, robust querying capabilities, and excellent data integrity, with a proven track record for mission-critical applications."
        
        elif 'mongodb' in tech_name_lower or 'nosql' in tech_name_lower:
            return "NoSQL database offering flexible schema design, horizontal scalability, and high performance for document-oriented data, ideal for rapid development and evolving data models."
        
        elif 'redis' in tech_name_lower or 'cache' in tech_name_lower:
            return "In-memory data structure store providing high-performance caching, pub/sub messaging, and data structure support, dramatically improving application response times."
        
        elif 'kubernetes' in tech_name_lower or 'k8s' in tech_name_lower:
            return "Kubernetes provides robust container orchestration with automated deployment, scaling, and management capabilities, ensuring reliable operation of containerized applications."
        
        elif 'docker' in tech_name_lower or 'container' in tech_name_lower:
            return "Containerization enables consistent environments across development and production, improving deployment reliability and simplifying dependency management."
        
        # Generic rationale based on category
        category_lower = category.lower()
        
        if 'frontend' in category_lower or 'ui' in category_lower:
            return f"{tech_name} was selected for its performance, development efficiency, and ability to create responsive user interfaces that meet modern user experience expectations."
        
        elif 'backend' in category_lower or 'server' in category_lower:
            return f"{tech_name} provides a robust server-side environment with excellent performance characteristics, security capabilities, and integration options."
        
        elif 'database' in category_lower or 'data' in category_lower:
            return f"{tech_name} offers the reliability, performance, and data management features required to support the application's data persistence needs."
        
        elif 'infrastructure' in category_lower or 'deploy' in category_lower:
            return f"{tech_name} enables reliable, scalable, and maintainable deployment of the application, supporting operational requirements for availability and performance."
        
        # Default rationale
        return f"{tech_name} was selected based on its capabilities to support {purpose}, considering factors like performance, reliability, and alignment with team expertise."
    
    def _extract_document_content(self, doc):
        """Extract structured content from Word document"""
        # This would use more advanced document parsing in a real implementation
        extracted_content = {
            'sections': {},
            'app_name': '',
            'description': ''
        }
        
        current_section = None
        current_subsection = None
        
        # Extract content from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Check for heading styles
            if para.style.name.startswith('Heading 1'):
                current_section = text
                current_subsection = None
                extracted_content['sections'][current_section] = {
                    'text': '',
                    'subsections': {}
                }
                
            elif para.style.name.startswith('Heading 2') and current_section:
                current_subsection = text
                extracted_content['sections'][current_section]['subsections'][current_subsection] = ''
                
            elif current_subsection and current_section:
                # Add to current subsection
                existing = extracted_content['sections'][current_section]['subsections'][current_subsection]
                extracted_content['sections'][current_section]['subsections'][current_subsection] = f"{existing}\n{text}" if existing else text
                
            elif current_section:
                # Add to current section
                existing = extracted_content['sections'][current_section]['text']
                extracted_content['sections'][current_section]['text'] = f"{existing}\n{text}" if existing else text
        
        # Extract basic info
        if 'Introduction' in extracted_content['sections']:
            extracted_content['description'] = extracted_content['sections']['Introduction']['text']
        
        # Try to find app name from title or content
        for para in doc.paragraphs:
            if para.style.name.startswith('Title') or para.style.name.startswith('Heading'):
                if para.text and 'HLDD' in para.text:
                    parts = para.text.split('HLDD')
                    if parts[0].strip():
                        extracted_content['app_name'] = parts[0].strip()
                        break
        
        return extracted_content
    
    def _enhance_extracted_content(self, extracted_content, enhancement_options):
        """Enhance extracted document content based on options"""
        enhanced_content = extracted_content.copy()
        
        # Enhance architecture section
        if enhancement_options.get('expand_architecture', False) and 'Architecture Overview' in extracted_content['sections']:
            arch_section = extracted_content['sections']['Architecture Overview']
            enhanced_arch = self._enhance_architecture_section(arch_section, extracted_content['description'])
            enhanced_content['sections']['Architecture Overview'] = enhanced_arch
        
        # Add or enhance security details
        if enhancement_options.get('add_security_details', False):
            if 'Security Architecture' not in extracted_content['sections']:
                # Create security section if missing
                enhanced_content['sections']['Security Architecture'] = {
                    'text': 'Security architecture for the system.',
                    'subsections': {
                        'Authentication and Authorization': self._generate_security_content('auth'),
                        'Data Protection': self._generate_security_content('data'),
                        'Network Security': self._generate_security_content('network'),
                        'Compliance': self._generate_security_content('compliance')
                    }
                }
            else:
                # Enhance existing security section
                security_section = extracted_content['sections']['Security Architecture']
                enhanced_security = self._enhance_security_section(security_section)
                enhanced_content['sections']['Security Architecture'] = enhanced_security
        
        # Add technology details
        if enhancement_options.get('update_technology', False) and 'Technology Stack' in extracted_content['sections']:
            tech_section = extracted_content['sections']['Technology Stack']
            enhanced_tech = self._enhance_technology_section(tech_section)
            enhanced_content['sections']['Technology Stack'] = enhanced_tech
        
        # Add best practices
        if enhancement_options.get('add_best_practices', False):
            if 'Architecture Principles' not in extracted_content['sections']:
                # Create principles section if missing
                enhanced_content['sections']['Architecture Principles'] = {
                    'text': 'Key principles guiding the architecture design:',
                    'subsections': {
                        'Design Principles': self._generate_principles_content('design'),
                        'Scalability Principles': self._generate_principles_content('scalability'),
                        'Security Principles': self._generate_principles_content('security'),
                        'Maintainability Principles': self._generate_principles_content('maintainability')
                    }
                }
        
        return enhanced_content
    
    def _enhance_architecture_section(self, arch_section, description):
        """Enhance architecture section content"""
        enhanced = arch_section.copy()
        
        # Expand main text if minimal
        if len(enhanced['text'].split()) < 50:
            enhanced['text'] = self._generate_expanded_architecture_overview(enhanced['text'], description)
        
        # Add common subsections if missing
        if 'Architecture Patterns' not in enhanced['subsections']:
            enhanced['subsections']['Architecture Patterns'] = self._generate_architecture_patterns_content(description)
        
        if 'Component Interactions' not in enhanced['subsections']:
            enhanced['subsections']['Component Interactions'] = self._generate_component_interactions_content(enhanced['subsections'])
        
        if 'Quality Attributes' not in enhanced['subsections']:
            enhanced['subsections']['Quality Attributes'] = self._generate_quality_attributes_content()
        
        return enhanced
    
    def _enhance_security_section(self, security_section):
        """Enhance security section content"""
        enhanced = security_section.copy()
        
        # Expand main text if minimal
        if len(enhanced['text'].split()) < 50:
            enhanced['text'] = (
                "The security architecture addresses protection of the system and its data "
                "through multiple layers of defense, following industry best practices and "
                "compliance requirements. It covers authentication, authorization, data protection, "
                "and network security to provide comprehensive protection against threats."
            )
        
        # Add common subsections if missing
        if 'Threat Model' not in enhanced['subsections']:
            enhanced['subsections']['Threat Model'] = self._generate_security_content('threat_model')
        
        if 'Security Controls' not in enhanced['subsections']:
            enhanced['subsections']['Security Controls'] = self._generate_security_content('controls')
        
        return enhanced
    
    def _enhance_technology_section(self, tech_section):
        """Enhance technology section content"""
        enhanced = tech_section.copy()
        
        # Expand main text if minimal
        if len(enhanced['text'].split()) < 50:
            enhanced['text'] = (
                "The technology stack has been selected to provide a robust, scalable, and maintainable "
                "foundation for the application. Each technology choice was made considering factors "
                "such as performance, reliability, team expertise, integration capabilities, and "
                "alignment with architectural goals. The stack is organized into layers corresponding "
                "to different aspects of the system."
            )
        
        # Add rationale subsection if missing
        if 'Technology Selection Rationale' not in enhanced['subsections']:
            enhanced['subsections']['Technology Selection Rationale'] = (
                "Technologies were selected based on the following criteria:\n"
                "- Alignment with architectural requirements and quality attributes\n"
                "- Performance and scalability characteristics\n"
                "- Maturity and community support\n"
                "- Security track record and features\n"
                "- Integration capabilities with other components\n"
                "- Team expertise and learning curve\n"
                "- Total cost of ownership including licensing and operational costs"
            )
        
        return enhanced
    
    def _generate_expanded_architecture_overview(self, current_text, description):
        """Generate expanded architecture overview text"""
        # Start with current text and expand
        expanded = current_text + "\n\n" if current_text else ""
        
        # Add general architecture description
        expanded += (
            "The architecture is designed to support the system's functional and non-functional "
            "requirements, providing a foundation for scalability, reliability, security, and "
            "maintainability. It follows industry best practices and patterns to address the "
            "specific challenges of the application domain.\n\n"
            "Components are organized with clear responsibilities and interfaces, enabling "
            "independent development, testing, and evolution. The architecture establishes "
            "boundaries that improve team autonomy while ensuring overall system coherence."
        )
        
        return expanded
    
    def _generate_architecture_patterns_content(self, description):
        """Generate content about architecture patterns"""
        # This would use more advanced NLP in a real implementation
        
        # Identify likely patterns from description
        description_lower = description.lower()
        
        patterns = []
        
        if any(term in description_lower for term in ['web', 'user interface', 'ui', 'frontend']):
            patterns.append({
                'name': 'Layered Architecture',
                'description': (
                    "Organizes the system into layers (presentation, business logic, data access) "
                    "with dependencies flowing downward. Each layer has a specific responsibility "
                    "and provides services to the layer above it."
                )
            })
        
        if any(term in description_lower for term in ['scalable', 'distributed', 'cloud']):
            patterns.append({
                'name': 'Microservices Architecture',
                'description': (
                    "Decomposes the application into small, independent services that communicate "
                    "over a network. Each service focuses on a specific business capability and "
                    "can be developed, deployed, and scaled independently."
                )
            })
        
        if any(term in description_lower for term in ['event', 'message', 'queue', 'stream']):
            patterns.append({
                'name': 'Event-Driven Architecture',
                'description': (
                    "Uses events to trigger and communicate between decoupled services. Components "
                    "emit events when state changes occur, and interested components consume those "
                    "events, enabling loose coupling and scalability."
                )
            })
        
        # Ensure we have at least one pattern
        if not patterns:
            patterns.append({
                'name': 'Component-Based Architecture',
'description': (
                    "Decomposes the system into independent, reusable components that expose well-defined interfaces. "
                    "Components can be developed and tested independently, promoting reusability and maintainability."
                )
            })
        
        # Format the content
        content = "The architecture employs the following patterns:\n\n"
        for pattern in patterns:
            content += f"**{pattern['name']}**: {pattern['description']}\n\n"
        
        return content
    
    def _generate_component_interactions_content(self, subsections):
        """Generate content about component interactions"""
        # Generate based on existing component information if available
        components = []
        
        # Look for component descriptions in subsections
        for subsection_name, subsection_content in subsections.items():
            if 'component' in subsection_name.lower():
                components.append(subsection_name)
        
        # If no components found, use generic content
        if not components:
            return (
                "Components interact through well-defined interfaces, primarily using synchronous REST APIs "
                "and asynchronous messaging. Communication patterns are selected based on requirements for "
                "coupling, consistency, and responsiveness.\n\n"
                "Key interactions include:\n"
                "- User Interface to API Gateway: Synchronous REST API calls\n"
                "- API Gateway to Services: Synchronous REST API calls\n"
                "- Inter-service communication: Asynchronous messaging for eventual consistency\n"
                "- Services to Data Layer: Direct database access or data access service APIs"
            )
        
        # Generate specific content based on found components
        content = (
            "Components interact through well-defined interfaces, using appropriate communication "
            "patterns based on requirements for coupling, consistency, and responsiveness.\n\n"
            "Key interactions include:\n"
        )
        
        # Generate interactions between components
        for i in range(len(components)):
            if i < len(components) - 1:
                content += f"- {components[i]} to {components[i+1]}: Synchronous API calls for direct operations\n"
            
            # Add data store interaction for last component
            if i == len(components) - 1:
                content += f"- {components[i]} to Data Store: Data access operations via repository pattern\n"
        
        return content
    
    def _generate_quality_attributes_content(self):
        """Generate content about quality attributes"""
        return (
            "The architecture addresses the following quality attributes:\n\n"
            "**Scalability**:\n"
            "- Horizontal scaling of stateless components\n"
            "- Caching strategies for performance optimization\n"
            "- Database sharding for data tier scalability\n\n"
            "**Reliability**:\n"
            "- Redundancy for critical components\n"
            "- Circuit breakers for fault isolation\n"
            "- Graceful degradation under load\n\n"
            "**Security**:\n"
            "- Defense in depth strategy\n"
            "- Least privilege principle\n"
            "- Encryption for data protection\n\n"
            "**Maintainability**:\n"
            "- Clear separation of concerns\n"
            "- Consistent coding standards\n"
            "- Comprehensive documentation"
        )
    
    def _generate_security_content(self, aspect):
        """Generate security content for a specific aspect"""
        if aspect == 'auth':
            return (
                "The system implements a robust authentication and authorization framework:\n\n"
                "- Authentication: OAuth 2.0 / OpenID Connect with JWT tokens\n"
                "- Multi-factor authentication for sensitive operations\n"
                "- Role-based access control (RBAC) for authorization\n"
                "- Centralized identity management\n"
                "- Session management with secure token handling"
            )
        elif aspect == 'data':
            return (
                "Data protection measures ensure confidentiality and integrity:\n\n"
                "- Encryption for data in transit (TLS 1.3)\n"
                "- Encryption for sensitive data at rest (AES-256)\n"
                "- Data masking for sensitive information in logs\n"
                "- Secure key management\n"
                "- Data classification and handling procedures"
            )
        elif aspect == 'network':
            return (
                "Network security controls protect the system boundary:\n\n"
                "- Web Application Firewall (WAF) for common attack protection\n"
                "- Network segmentation with security groups\n"
                "- Rate limiting and throttling\n"
                "- DDoS protection\n"
                "- Intrusion detection and prevention systems"
            )
        elif aspect == 'compliance':
            return (
                "The system is designed to meet relevant compliance requirements:\n\n"
                "- GDPR compliance for data privacy\n"
                "- Security logging and audit trails\n"
                "- Regular security assessments\n"
                "- Vulnerability management process\n"
                "- Security incident response procedures"
            )
        elif aspect == 'threat_model':
            return (
                "The security architecture addresses the following threat categories:\n\n"
                "- Authentication bypassing and session hijacking\n"
                "- Authorization bypassing and privilege escalation\n"
                "- Injection attacks (SQL, NoSQL, command injection, etc.)\n"
                "- Cross-site scripting and cross-site request forgery\n"
                "- Sensitive data exposure\n"
                "- Denial of service attacks\n\n"
                "Each threat category has been analyzed with associated mitigations implemented."
            )
        elif aspect == 'controls':
            return (
                "Security controls are implemented at multiple layers:\n\n"
                "**Application Layer Controls**:\n"
                "- Input validation and output encoding\n"
                "- Authentication and authorization\n"
                "- Secure session management\n"
                "- Secure configuration and hardening\n\n"
                "**Infrastructure Layer Controls**:\n"
                "- Network segmentation and firewalls\n"
                "- Intrusion detection/prevention systems\n"
                "- Secure configuration and patching\n"
                "- Monitoring and logging\n\n"
                "**Operational Controls**:\n"
                "- Security incident response\n"
                "- Vulnerability management\n"
                "- Change management\n"
                "- Security awareness training"
            )
        else:
            return "Security content for this aspect is not available."
    
    def _generate_principles_content(self, principle_type):
        """Generate content about architecture principles"""
        if principle_type == 'design':
            return (
                "**Separation of Concerns**: Divide the system into distinct features with minimal overlap\n\n"
                "**Single Responsibility Principle**: Each component should have one reason to change\n\n"
                "**Don't Repeat Yourself (DRY)**: Avoid duplication by abstracting common functionality\n\n"
                "**Interface Segregation**: Clients should not depend on interfaces they don't use\n\n"
                "**Dependency Inversion**: High-level modules should not depend on low-level modules"
            )
        elif principle_type == 'scalability':
            return (
                "**Horizontal Scaling**: Design components to scale out rather than up\n\n"
                "**Statelessness**: Minimize state management in application components\n\n"
                "**Asynchronous Processing**: Use message queues for non-critical operations\n\n"
                "**Caching Strategy**: Implement appropriate caching at multiple levels\n\n"
                "**Data Partitioning**: Design for data sharding and partitioning from the start"
            )
        elif principle_type == 'security':
            return (
                "**Defense in Depth**: Implement security at multiple layers\n\n"
                "**Least Privilege**: Provide only the minimum necessary access\n\n"
                "**Secure by Default**: Systems should be secure in their default configuration\n\n"
                "**Fail Securely**: Errors should not compromise security\n\n"
                "**Security as a Process**: Security is continuous, not a one-time effort"
            )
        elif principle_type == 'maintainability':
            return (
                "**Simplicity**: Avoid unnecessary complexity\n\n"
                "**Consistency**: Follow consistent patterns and conventions\n\n"
                "**Automation**: Automate build, test, and deployment processes\n\n"
                "**Testability**: Design for testability at all levels\n\n"
                "**Documentation**: Maintain appropriate documentation for the architecture"
            )
        else:
            return "Principles for this category are not available."
    
    def _create_document_from_content(self, content):
        """Create a Word document from extracted and enhanced content"""
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add document properties
        core_properties = doc.core_properties
        core_properties.title = f"{content.get('app_name', 'Application')} High Level Design Document"
        core_properties.subject = "Architecture Documentation"
        core_properties.creator = "Architecture Team"
        
        # Add cover page
        app_name = content.get('app_name', 'Application')
        self._add_cover_page(doc, {'name': app_name})
        
        # Add table of contents
        self._add_toc(doc)
        
        # Add sections in logical order
        section_order = [
            'Introduction',
            'System Overview',
            'Architecture Overview',
            'Architecture Principles',
            'Technology Stack',
            'Data Architecture',
            'Security Architecture',
            'Deployment Architecture',
            'Operational Considerations',
            'Appendices'
        ]
        
        # Add each section with its content
        for section_name in section_order:
            if section_name in content['sections']:
                section_data = content['sections'][section_name]
                
                # Add section heading
                doc.add_heading(section_name, level=1)
                
                # Add section text
                if section_data['text']:
                    for paragraph in section_data['text'].split('\n'):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
                
                # Add subsections
                for subsection_name, subsection_text in section_data['subsections'].items():
                    doc.add_heading(subsection_name, level=2)
                    
                    # Handle markdown-like formatting in text
                    lines = subsection_text.split('\n')
                    current_list = None
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Check for lists
                        if line.startswith('- '):
                            # Add bullet point
                            if current_list != 'bullet':
                                current_list = 'bullet'
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(line[2:])
                        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
                            # Handle bold headers
                            text = line.strip('**')
                            p = doc.add_paragraph()
                            p.add_run(text).bold = True
                        else:
                            # Regular paragraph
                            current_list = None
                            doc.add_paragraph(line)
        
        return doc
    
    def _generate_enhancement_summary(self, original_content, enhanced_content):
        """Generate a summary of enhancements made to the document"""
        summary = {
            'sections_added': [],
            'sections_enhanced': [],
            'subsections_added': [],
            'enhancement_count': 0
        }
        
        # Check for added sections
        for section_name in enhanced_content['sections']:
            if section_name not in original_content['sections']:
                summary['sections_added'].append(section_name)
                summary['enhancement_count'] += 1
        
        # Check for enhanced sections and added subsections
        for section_name, section_data in enhanced_content['sections'].items():
            if section_name in original_content['sections']:
                original_section = original_content['sections'][section_name]
                
                # Check if section text was enhanced
                original_text_length = len(original_section['text'].split())
                enhanced_text_length = len(section_data['text'].split())
                
                if enhanced_text_length > original_text_length * 1.2:  # 20% more content
                    summary['sections_enhanced'].append(section_name)
                    summary['enhancement_count'] += 1
                
                # Check for added subsections
                for subsection_name in section_data['subsections']:
                    if subsection_name not in original_section['subsections']:
                        summary['subsections_added'].append(f"{section_name} > {subsection_name}")
                        summary['enhancement_count'] += 1
        
        return summary
    
    def _load_patterns_knowledge_base(self):
        """Load architecture patterns knowledge base"""
        # In a real implementation, this would load from a database or file
        return {
            'microservices': {
                'description': 'Architecture with small, independent services that communicate over a network',
                'best_for': ['scalability', 'maintainability', 'agility'],
                'caution_for': ['cost', 'complexity', 'operational overhead'],
                'components': [
                    'API Gateway', 'Service Discovery', 'Container Orchestration', 
                    'Distributed Logging', 'Monitoring', 'CI/CD Pipeline'
                ]
            },
            'serverless': {
                'description': 'Event-driven architecture where code runs in stateless containers',
                'best_for': ['cost optimization', 'scalability', 'time_to_market'],
                'caution_for': ['vendor lock-in', 'cold starts', 'debugging'],
                'components': [
                    'Function-as-a-Service', 'API Gateway', 'Event Bus', 
                    'Managed Database', 'Storage', 'Monitoring'
                ]
            },
            'layered': {
                'description': 'Architecture organized in horizontal layers (presentation, business, data, etc.)',
                'best_for': ['maintainability', 'separation of concerns', 'testing'],
                'caution_for': ['performance', 'scalability', 'tight coupling'],
                'components': [
                    'Presentation Layer', 'Business Logic Layer', 'Data Access Layer', 
                    'Database', 'Cross-cutting Concerns'
                ]
            }
            # Additional patterns would be included here
        }
    
    def _load_technologies_knowledge_base(self):
        """Load technologies knowledge base"""
        # In a real implementation, this would load from a database or file
        return {
            'frontend': {
                'react': {
                    'description': 'JavaScript library for building user interfaces',
                    'strengths': ['Component-based', 'Virtual DOM', 'Large ecosystem'],
                    'weaknesses': ['Learning curve', 'Requires additional libraries for full framework']
                },
                'angular': {
                    'description': 'Platform and framework for building client applications',
                    'strengths': ['Full framework', 'TypeScript integration', 'Comprehensive'],
                    'weaknesses': ['Steeper learning curve', 'More opinionated']
                },
                'vue': {
                    'description': 'Progressive JavaScript framework for UIs',
                    'strengths': ['Gentle learning curve', 'Flexible', 'Good performance'],
                    'weaknesses': ['Smaller ecosystem than React/Angular']
                }
            },
            'backend': {
                'node': {
                    'description': 'JavaScript runtime for server-side applications',
                    'strengths': ['JavaScript across stack', 'Asynchronous I/O', 'Large ecosystem'],
                    'weaknesses': ['CPU-intensive tasks', 'Callback complexity']
                },
                'python': {
                    'description': 'High-level, interpreted programming language',
                    'strengths': ['Readability', 'Extensive libraries', 'Rapid development'],
                    'weaknesses': ['GIL for threading', 'Performance vs compiled languages']
                },
                'java': {
                    'description': 'General-purpose programming language',
                    'strengths': ['Performance', 'Enterprise-grade', 'Strong typing'],
                    'weaknesses': ['Verbosity', 'Development speed']
                }
            },
            'database': {
                'postgresql': {
                    'description': 'Advanced open-source relational database',
                    'strengths': ['ACID compliance', 'Advanced features', 'SQL standard conformance'],
                    'weaknesses': ['Horizontal scaling complexity']
                },
                'mongodb': {
                    'description': 'Document-oriented NoSQL database',
                    'strengths': ['Schema flexibility', 'Horizontal scaling', 'JSON-like documents'],
                    'weaknesses': ['Transaction model', 'Memory usage']
                },
                'redis': {
                    'description': 'In-memory data structure store',
                    'strengths': ['Performance', 'Versatility', 'Data structures'],
                    'weaknesses': ['Memory constraints', 'Persistence model']
                }
            }
            # Additional technologies would be included here
        }
    
    def _load_domain_knowledge_base(self):
        """Load domain-specific knowledge base"""
        # In a real implementation, this would load from a database or file
        return {
            'financial': {
                'compliance': ['PCI DSS', 'SOX', 'Basel III'],
                'security_requirements': ['Encryption at rest and in transit', 'Multi-factor authentication', 'Audit logging'],
                'key_components': ['Payment processing', 'Account management', 'Reporting', 'Fraud detection']
            },
            'healthcare': {
                'compliance': ['HIPAA', 'GDPR', 'HL7'],
                'security_requirements': ['PHI protection', 'Access controls', 'Audit trails'],
                'key_components': ['Patient records', 'Provider management', 'Billing', 'Clinical systems']
            },
            'e-commerce': {
                'compliance': ['PCI DSS', 'Consumer protection laws', 'GDPR'],
                'security_requirements': ['Secure payment processing', 'Data protection', 'Fraud prevention'],
                'key_components': ['Product catalog', 'Shopping cart', 'Payment processing', 'Order management']
            },
            'general': {
                'compliance': ['GDPR', 'Local data protection laws'],
                'security_requirements': ['Authentication', 'Authorization', 'Data protection'],
                'key_components': ['User management', 'Content management', 'Search functionality', 'Reporting']
            }
        }
    
    # The _create_document, _setup_document_styles, _add_cover_page, and _add_toc methods
    # can be copied from your existing DocumentGenerator class

    def _setup_document_styles(self, doc):
        """Set up document styles for consistent formatting"""
        # Heading styles
        for i in range(1, 5):
            style = doc.styles[f'Heading {i}']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(20 - (i * 2))  # Decreasing sizes: 18, 16, 14, 12
            font.bold = True
            font.color.rgb = None  # Black
        
        # Normal text style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create a style for the TOC title
        try:
            toc_title_style = doc.styles.add_style('TOC Title', WD_STYLE_TYPE.PARAGRAPH)
            toc_title_font = toc_title_style.font
            toc_title_font.name = 'Calibri'
            toc_title_font.size = Pt(16)
            toc_title_font.bold = True
        except:
            pass  # Style might already exist
        
        # Create caption style
        try:
            if 'Caption' not in doc.styles:
                caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
                caption_font = caption_style.font
                caption_font.name = 'Calibri'
                caption_font.size = Pt(10)
                caption_font.italic = True
                
                paragraph_format = caption_style.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Non-critical, continue without caption style
    
    def _add_cover_page(self, doc, application_info):
        """Add a cover page to the document"""
        # Add large title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"{application_info.get('name', 'Application')}")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Add subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("High Level Design Document")
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.bold = True
        
        # Add version info
        doc.add_paragraph()  # Spacing
        version_para = doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_para.add_run(f"Version: {application_info.get('version', '1.0')}")
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
        
        # Add author
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"Author: {application_info.get('author', 'Architecture Team')}")
        
        # Add a page break
        doc.add_page_break()
    
    def _add_toc(self, doc):
        """Add table of contents"""
        toc_para = doc.add_paragraph()
        try:
            toc_para.style = 'TOC Title'
        except:
            pass  # Style might not exist, continue without styling
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_para.add_run("Table of Contents")
        
        doc.add_paragraph()  # Spacing
        
        # Insert TOC field code
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # TOC field code
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._element.append(fldChar)
        
        fldChar = OxmlElement('w:t')
        fldChar.text = "Right-click to update table of contents."
        run._element.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        # Add a page break
        doc.add_page_break()
    
    def _create_document(self, application_info, tech_stack_info, architecture_info):
        """Create the HLDD document with proper formatting"""
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add document properties
        core_properties = doc.core_properties
        core_properties.title = f"{application_info.get('name', 'Application')} High Level Design Document"
        core_properties.subject = "Architecture Documentation"
        core_properties.creator = application_info.get('author', 'Architecture Team')
        
        # Create cover page
        self._add_cover_page(doc, application_info)
        
        # Add table of contents
        self._add_toc(doc)
        
        # Add document sections - enhanced with additional sections and better content
        self._add_introduction(doc, application_info)
        self._add_system_overview(doc, application_info)
        self._add_architecture_overview(doc, architecture_info)
        self._add_architecture_principles(doc, architecture_info)
        self._add_technology_stack(doc, tech_stack_info)
        self._add_data_architecture(doc, architecture_info.get('data_architecture', {}))
        self._add_security_architecture(doc, architecture_info.get('security', {}))
        self._add_deployment_architecture(doc, architecture_info.get('deployment', {}))
        self._add_operational_considerations(doc, architecture_info.get('operations', {}))
        self._add_appendices(doc, application_info)
        
        return doc
    
    # The remainder of the document section generation methods (_add_introduction, etc.)
    # can be copied from your existing DocumentGenerator class, with enhancements
    # as needed to support the AI-driven content generation

    def _add_introduction(self, doc, application_info):
        """Add introduction section with enhanced content"""
        doc.add_heading('Introduction', level=1)
        
        doc.add_heading('Purpose', level=2)
        purpose_text = application_info.get('purpose', 'This High Level Design Document (HLDD) provides an overview of the system architecture, describing the structural components, their relationships, and the design decisions that went into creating the architecture.')
        doc.add_paragraph(purpose_text)
        
        doc.add_heading('Scope', level=2)
        scope_text = application_info.get('scope', 'This document describes the high-level design of the system including its components, interfaces, data structures, and technologies. It serves as a reference for stakeholders including developers, architects, and operations teams. It provides sufficient detail to understand the overall system structure without delving into low-level implementation details.')
        doc.add_paragraph(scope_text)
        
        doc.add_heading('Intended Audience', level=2)
        audience_text = application_info.get('audience', 'This document is intended for:\n- Development teams implementing the system\n- Architecture teams evaluating or extending the system\n- Operations teams responsible for deployment and maintenance\n- Technical stakeholders requiring system understanding')
        
        # Handle multi-line content with bullet points
        if '\n' in audience_text:
            lines = audience_text.split('\n')
            doc.add_paragraph(lines[0])
            for line in lines[1:]:
                if line.strip().startswith('- '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line.strip()[2:])
                else:
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph(audience_text)
        
        doc.add_heading('Definitions, Acronyms, and Abbreviations', level=2)
        
        # Add a table for definitions - with enhanced content
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Term/Acronym'
        header_cells[1].text = 'Definition'
        
        # Make the header row bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add enhanced definitions based on content
        definitions = application_info.get('definitions', [
            {'term': 'HLDD', 'definition': 'High Level Design Document'},
            {'term': 'API', 'definition': 'Application Programming Interface - Defines how software components interact'},
            {'term': 'REST', 'definition': 'Representational State Transfer - Architectural style for distributed systems'},
            {'term': 'DB', 'definition': 'Database - Organized collection of structured data'},
            {'term': 'UI', 'definition': 'User Interface - Visual elements through which users interact with software'},
            {'term': 'SLA', 'definition': 'Service Level Agreement - Commitment between service provider and client'},
            {'term': 'CI/CD', 'definition': 'Continuous Integration/Continuous Deployment - Automated software delivery practices'},
            {'term': 'MFA', 'definition': 'Multi-Factor Authentication - Multiple methods to verify identity'},
            {'term': 'TLS', 'definition': 'Transport Layer Security - Cryptographic protocol for secure communications'}
        ])
        
        for definition in definitions:
            row = table.add_row()
            row.cells[0].text = definition.get('term', '')
            row.cells[1].text = definition.get('definition', '')

# Factory function to get an instance of the Enhanced Document Generator
def get_enhanced_document_generator(output_dir=None):
    """Factory function to get an instance of the Enhanced Document Generator"""
    return EnhancedDocumentGenerator(output_dir)