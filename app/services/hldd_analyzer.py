# app/services/hldd_analyzer.py
# -*- coding: utf-8 -*-

import os
import re
import docx
import fitz  # PyMuPDF
import pandas as pd
import numpy as np
import json
import logging
from app.config import Config
from textblob import TextBlob  # For basic NLP tasks
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify, current_app
from flask_login import login_required, current_user
from app.extensions import db

from app.models.application import Application
from app.models.arb_review import ARBReview
from werkzeug.utils import secure_filename


class HLDDAnalyzer:
    """
    Service for analyzing High Level Design Documents (HLDD) and evaluating
    architectural soundness
    """
    
    def __init__(self, hldd_path):
        """
        Initializes the class with the path to the High-Level Design Document (HLDD) and sets up
        data structures for architecture analysis.

        Args:
            hldd_path (str): The file path to the HLDD document.
        """

        self.hldd_path = hldd_path
        self.hldd_content = None  # Will store the loaded HLDD content
        self.architecture_patterns = {
            'microservices': {'weight': 0.10},
            'layered': {'weight': 0.15},
            'event-driven': {'weight': 0.12},
            'serverless': {'weight': 0.08},
            'service-oriented': {'weight': 0.10},
            'monolithic': {'weight': 0.05},
            'client-server': {'weight': 0.05},
            'peer-to-peer': {'weight': 0.03},
            'space-based': {'weight': 0.02},
            'api gateway': {'weight': 0.08},
            'circuit breaker': {'weight': 0.07},
            'cqrs': {'weight': 0.05}
        }
        self.anti_patterns = {
            'big ball of mud': {'weight': 0.15},
            'spaghetti code': {'weight': 0.12},
            'monolith': {'weight': 0.10},
            'tight coupling': {'weight': 0.20},
            'god class': {'weight': 0.08},
            'silver bullet': {'weight': 0.05},
            'golden hammer': {'weight': 0.07},
            'distributed monolith': {'weight': 0.13},
            'analysis paralysis': {'weight': 0.03},
            'vendor lock-in': {'weight': 0.05},
            'reinventing the wheel': {'weight': 0.02}
        }
        self.quality_attributes = {
            'scalability': {'weight': 0.20},
            'performance': {'weight': 0.20},
            'reliability': {'weight': 0.15},
            'availability': {'weight': 0.15},
            'maintainability': {'weight': 0.10},
            'security': {'weight': 0.10},
            'usability': {'weight': 0.05},
            'accessibility': {'weight': 0.05}
        }
        
        self.criteria = {
            'security': {
                'weight': 0.25,
                'keywords': [
                    'encryption', 'authentication', 'authorization',
                    'identity management', 'access control lists (ACLs)', 'role-based access control (RBAC)',
                    'least privilege', 'multi-factor authentication (MFA)', 'single sign-on (SSO)',
                    'data at rest encryption', 'data in transit encryption', 'end-to-end encryption',
                    'transport layer security (TLS)', 'secure sockets layer (SSL)', 'certificates',
                    'firewall', 'web application firewall (WAF)', 'intrusion detection system (IDS)',
                    'intrusion prevention system (IPS)', 'vulnerability scanning', 'penetration testing',
                    'threat modeling', 'security audit', 'compliance frameworks', 'data breach prevention',
                    'incident response', 'security policies', 'security best practices', 'zero trust',
                    'data masking', 'tokenization', 'hashing', 'salting', 'key management'
                ]
            },
            'scalability': {
                'weight': 0.20,
                'keywords': [
                    'scaling up', 'scaling out', 'horizontal scaling', 'vertical scaling', 'auto-scaling groups',
                    'load balancing', 'distributed systems', 'distributed databases', 'sharding', 'partitioning',
                    'replication', 'caching strategies', 'content delivery network (CDN)', 'message queues',
                    'event-driven architecture', 'microservices architecture', 'containerization (Docker, Podman)',
                    'orchestration (Kubernetes, ECS, Swarm)', 'elasticity', 'throughput optimization',
                    'performance tuning', 'capacity planning', 'high-performance computing (HPC)',
                    'reactive programming', 'non-blocking I/O', 'concurrency control', 'stateless services'
                ]
            },
            'reliability': {
                'weight': 0.20,
                'keywords': [
                    'high availability (HA)', 'fault tolerance', 'redundancy (active-active, active-passive)',
                    'failover mechanisms', 'automatic recovery', 'disaster recovery (DR)', 'business continuity (BC)',
                    'backup and restore', 'data replication', 'data integrity checks', 'checksums',
                    'circuit breakers', 'retry mechanisms', 'idempotency', 'monitoring and alerting',
                    'health checks', 'service level agreements (SLAs)', 'uptime guarantees', 'mean time between failures (MTBF)',
                    'mean time to recovery (MTTR)', 'graceful degradation', 'self-healing', 'observability (metrics, logs, traces)'
                ]
            },
            'maintainability': {
                'weight': 0.15,
                'keywords': [
                    'documentation (design, API, user)', 'code readability', 'code modularity', 'low coupling', 'high cohesion',
                    'naming conventions', 'code comments', 'architectural patterns', 'design principles (SOLID, DRY, KISS)',
                    'refactoring', 'technical debt management', 'version control (Git, SVN)', 'branching strategies',
                    'continuous integration/continuous delivery (CI/CD) pipelines', 'automated testing (unit, integration, E2E)',
                    'infrastructure as code (IaC) (Terraform, CloudFormation)', 'configuration management (Ansible, Chef, Puppet)',
                    'monitoring dashboards', 'logging standards', 'centralized logging', 'distributed tracing',
                    'observability platforms', 'dependency management', 'code reviews', 'static code analysis', 'linting'
                ]
            },
            'cost_optimization': {
                'weight': 0.10,
                'keywords': [
                    'cost-effectiveness', 'budget constraints', 'resource optimization', 'waste reduction',
                    'reserved instances', 'savings plans', 'spot instances', 'elastic compute', 'serverless computing (Lambda, Functions)',
                    'pay-as-you-go pricing', 'rightsizing resources', 'storage tiering', 'data lifecycle management',
                    'network optimization', 'caching to reduce costs', 'cost monitoring tools', 'cloud financial management (FinOps)',
                    'resource tagging', 'cost allocation', 'automation of resource management', 'economies of scale'
                ]
            },
            'compliance': {
                'weight': 0.10,
                'keywords': [
                    'General Data Protection Regulation (GDPR)', 'Health Insurance Portability and Accountability Act (HIPAA)',
                    'Payment Card Industry Data Security Standard (PCI DSS)', 'Sarbanes-Oxley Act (SOX)',
                    'Federal Risk and Authorization Management Program (FedRAMP)', 'regulatory compliance',
                    'legal requirements', 'industry standards', 'internal policies', 'governance frameworks',
                    'data protection laws', 'privacy regulations', 'data sovereignty', 'audit trails for compliance',
                    'access logging for compliance', 'security controls for compliance', 'risk and compliance management',
                    'data residency', 'data localization', 'compliance certifications'
                ]
            }
        }

        # Required architecture components
        self.required_components = {
            'database': {'weight': 0.20},
            'application server': {'weight': 0.20},
            'networking': {'weight': 0.15},
            'security': {'weight': 0.15},
            'authentication': {'weight': 0.10},
            'monitoring': {'weight': 0.10},
            'backup': {'weight': 0.05},
            'disaster recovery': {'weight': 0.05}
        }
        
        # Load the HLDD content
        self.load_hldd()
    
    def load_hldd(self):
        """Load the HLDD content from the specified path"""
        try:
            if not os.path.exists(self.hldd_path):
                logging.warning(f"HLDD file not found at path: {self.hldd_path}")
                self.hldd_content = ""
                return
                
            # Load the document content based on file extension
            _, file_extension = os.path.splitext(self.hldd_path)
            
            if file_extension.lower() == '.docx':
                self.hldd_content = self._read_docx(self.hldd_path)
            elif file_extension.lower() == '.pdf':
                self.hldd_content = self._read_pdf(self.hldd_path)
            elif file_extension.lower() in ['.txt', '.md']:
                self.hldd_content = self._read_text(self.hldd_path)
            else:
                logging.warning(f"Unsupported file format: {file_extension}")
                self.hldd_content = ""
                
        except Exception as e:
            logging.error(f"Error loading HLDD content: {str(e)}")
            self.hldd_content = ""  # Set empty content on error
    
    def read_document(self, file_path):
        """Extract text content from various document formats"""
        _, file_extension = os.path.splitext(file_path)
        
        if file_extension.lower() == '.docx':
            return self._read_docx(file_path)
        elif file_extension.lower() == '.pdf':
            return self._read_pdf(file_path)
        elif file_extension.lower() in ['.txt', '.md']:
            return self._read_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _read_docx(self, file_path):
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    
    def _read_pdf(self, file_path):
        """Extract text from PDF file"""
        doc = fitz.open(file_path)
        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        return '\n'.join(full_text)
    
    def _read_text(self, file_path):
        """Read plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def analyze_document(self, file_path=None):
        """
        Analyze HLDD document and return architecture evaluation results
        
        Args:
            file_path (str, optional): Path to the HLDD document. If not provided,
                                       uses the already loaded content.
        
        Returns:
            dict: Analysis results
        """
        try:
            # Extract text from document
            if file_path:
                document_text = self.read_document(file_path)
            else:
                # Use already loaded content
                if not self.hldd_content:
                    self.load_hldd()
                document_text = self.hldd_content
            
            # Check if we have content to analyze
            if not document_text:
                return {
                    'error': "No document content to analyze",
                    'scores': {},
                    'missing_components': [],
                    'overall_score': 0,
                    'acceptable': False,
                    'recommendations': ["Please check the document file and try again."]
                }
            
            # Run analysis
            results = {
                'scores': self._calculate_criteria_scores(document_text),
                'missing_components': self._identify_missing_components(document_text),
                'overall_score': 0,
                'acceptable': False,
                'recommendations': []
            }
            
            # Calculate overall weighted score
            overall_score = 0
            for criterion, data in results['scores'].items():
                overall_score += data['score'] * self.criteria[criterion]['weight']
            
            results['overall_score'] = round(overall_score, 2)
            results['acceptable'] = overall_score >= 0.7 and len(results['missing_components']) <= 1
            
            # Generate recommendations
            results['recommendations'] = self._generate_recommendations(results)
            
            return results
        
        except Exception as e:
            logging.error(f"Error analyzing document: {str(e)}")
            return {
                'error': str(e),
                'scores': {},
                'missing_components': [],
                'overall_score': 0,
                'acceptable': False,
                'recommendations': ["Error analyzing document. Please check file format and try again."]
            }
    
    def _calculate_criteria_scores(self, text):
        """Calculate scores for each evaluation criterion"""
        text = text.lower()
        scores = {}
        
        for criterion, data in self.criteria.items():
            keyword_count = 0
            for keyword in data['keywords']:
                keyword_count += len(re.findall(r'\b' + re.escape(keyword) + r'\b', text))
            
            # Normalize score (0-1 range)
            max_expected_occurrences = len(data['keywords']) * 2
            normalized_score = min(1.0, keyword_count / max_expected_occurrences)
            
            scores[criterion] = {
                'score': round(normalized_score, 2),
                'keywords_found': keyword_count
            }
        
        return scores
    
    def _identify_missing_components(self, text):
        """Identify any required components missing from the architecture"""
        text = text.lower()
        missing_components = []
        
        for component in self.required_components:
            if re.search(r'\b' + re.escape(component) + r'\b', text) is None:
                missing_components.append(component)
        
        return missing_components
    
    def _generate_recommendations(self, results):
        """Generate recommendations based on analysis results"""
        recommendations = []
        
        # Recommend addressing missing components
        if results['missing_components']:
            components_list = ', '.join(results['missing_components'])
            recommendations.append(
                f"Add details about the following components: {components_list}"
            )
        
        # Recommend improvements for low-scoring criteria
        for criterion, data in results['scores'].items():
            if data['score'] < 0.5:
                recommendations.append(
                    f"Enhance {criterion} details. Consider adding more information about " +
                    f"{', '.join(self.criteria[criterion]['keywords'][:3])}"
                )
        
        return recommendations


# Helper functions for document analysis

def get_analyzer():
    """
    Initialize and return an HLDD analyzer
    
    Returns:
        dict: A dictionary containing analyzer configuration or methods
    """
    try:
        # Example minimal implementation
        return {
            'analyze_docx': analyze_docx,
            'analyze_pdf': analyze_pdf,
            'extract_text': extract_document_text,
            'analyze_csv': analyze_csv
        }
    except Exception as e:
        logging.error(f"Error initializing HLDD analyzer: {e}")
        return {}

def analyze_docx(file_path):
    """
    Analyze a Word document for HLDD content
    
    Args:
        file_path (str): Path to the Word document
    
    Returns:
        dict: Analysis results
    """
    try:
        # Basic document analysis
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        return {
            'success': True,
            'paragraphs': paragraphs,
            'paragraph_count': len(paragraphs)
        }
    except Exception as e:
        logging.error(f"Error analyzing DOCX document: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def analyze_pdf(file_path):
    """
    Analyze a PDF document for HLDD content
    
    Args:
        file_path (str): Path to the PDF document
    
    Returns:
        dict: Analysis results
    """
    try:
        # Open the PDF
        pdf_document = fitz.open(file_path)
        
        # Extract text from all pages
        text_pages = []
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text_pages.append(page.get_text())
        
        # Combine texts
        full_text = '\n'.join(text_pages)
        
        return {
            'success': True,
            'pages': text_pages,
            'page_count': len(text_pages),
            'total_text': full_text
        }
    except Exception as e:
        logging.error(f"Error analyzing PDF document: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def analyze_csv(file_path):
    """
    Analyze a CSV file using pandas
    
    Args:
        file_path (str): Path to the CSV file
    
    Returns:
        dict: Analysis results
    """
    try:
        # Read CSV file
        df = pd.read_csv(file_path)
        
        return {
            'success': True,
            'columns': list(df.columns),
            'row_count': len(df),
            'column_count': len(df.columns),
            'summary': df.describe().to_dict()
        }
    except Exception as e:
        logging.error(f"Error analyzing CSV document: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def extract_document_text(file_path):
    """
    Extract full text from a document based on its extension
    
    Args:
        file_path (str): Path to the document
    
    Returns:
        str: Extracted text from the document
    """
    try:
        # Determine file type
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.docx':
            doc = docx.Document(file_path)
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == '.pdf':
            pdf_document = fitz.open(file_path)
            text_pages = [page.get_text() for page in pdf_document]
            full_text = '\n'.join(text_pages)
        elif ext == '.csv':
            df = pd.read_csv(file_path)
            full_text = df.to_string()
        else:
            return f"Unsupported file type: {ext}"
        
        return full_text
    except Exception as e:
        logging.error(f"Error extracting document text: {e}")
        return ""